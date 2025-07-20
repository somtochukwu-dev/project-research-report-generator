import os
from docx import Document
from docx.shared import Inches
from datetime import datetime


def create_report(template_path, output_path, stats_df, figure_paths, metadata):
    """
    Build a research report using a Word template and export it.

    Parameters:
        template_path (str): Path to the .docx template file
        output_path (str): Destination path for the final report
        stats_df (pd.DataFrame): Summary statistics to include
        figure_paths (list): List of image file paths to embed
        metadata (dict): Dictionary of report metadata (title, date, author, etc.)
    """

    # Load or create a document
    if os.path.exists(template_path):
        doc = Document(template_path)
    else:
        doc = Document()

    # Title Section
    doc.add_heading(metadata.get("title", "Experiment Report"), level=1)
    doc.add_paragraph(f"Author: {metadata.get('author', 'N/A')}")
    doc.add_paragraph(f"Date: {metadata.get('date', datetime.today().strftime('%Y-%m-%d'))}")
    doc.add_paragraph("")

    doc.add_heading("Overview", level=2)
    doc.add_paragraph(metadata.get("description", "Automated report generated from Sci-Cle experimental log processor."))

    # Summary Stats
    doc.add_heading("Summary Statistics", level=2)
    table = doc.add_table(rows=1, cols=len(stats_df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(stats_df.columns):
        hdr_cells[i].text = str(col)

    for _, row in stats_df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(round(val, 3)) if isinstance(val, float) else str(val)

    # Visualisations
    doc.add_heading("Visual Analysis", level=2)
    for fig_path in figure_paths:
        doc.add_paragraph(os.path.basename(fig_path).replace('_', ' ').replace('.png', '').title())
        doc.add_picture(fig_path, width=Inches(5.5))
        doc.add_paragraph("")

    # Footer
    doc.add_paragraph("Report generated automatically using Sci-Cle.")
    doc.add_paragraph("© 2025 Sci-Cle Automation – Precision for Research.")

    # Export
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"📄 Report saved to: {output_path}")
